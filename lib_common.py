
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from bs4 import BeautifulSoup
from docx import Document

import re

def split_in_paragraphs(text):
    
    def lookup_biblice():
        vechiul_testament = ["geneza", "facere", "exodul", "iesirea", "levitic", "numerii", "deuteronomul", "iosua", "judecatori", "rut", "regi1", "1regi", "regi2", "2regi", "regi3", "3regi", "regi4", "4regi", "cronici1", "paralelipomena1", "1paralelipomena", "cronici2", "paralelipomena2", "2paralelipomena", "ezdra", "neemia", "estera", "iov", "psalmi", "proverbe", "ecclesiastul", "cantarea cantarilor", "isaia", "ieremia", "plangerile ieremia", "iezechiel", "daniel", "osea", "amos", "miheia", "ioil", "avdie", "iona", "naum", "avacum", "sofonie", "agheu", "zaharia", "maleahi"] # ar mai fi de completat

        noul_testament = ["matei", "marcu", "luca", "ioan", "faptele apostolilor", "romani", "corinteni1", "1corinteni", "corinteni2", "2corinteni", "galateni", "efeseni", "filipeni", "coloseni", "tesaloniceni1", "tesaloniceni2", "timotei1", "1timotei", "timotei2", "2timotei", "tit", "filimon", "evrei", "iacov", "petru1", "1petru", "petru2", "2petru", "ioan1", "1ioan", "ioan2", "2ioan", "ioan3", "3ioan", "iuda", "apocalipsa"]
        
        def genereaza_combinatiile(lista):
            comb_lista = []
            vocale = set('aeiouAEIOU')  # setul de vocale
            for cuvant in lista:
                for i in range(2, min(len(cuvant), 5) + 1):  
                    comb_lista.append(cuvant[:i])

                # cautam prima consoana dupa prima litera si o adaugam la lista
                for litera in cuvant[1:]:  # pornim de la a doua litera
                    if litera not in vocale:
                        comb_lista.append(cuvant[0] + litera)
                        #break  # ne oprim dupa ce am gasit prima consoana -, nu, continuam ! fiindca la marcu, avem mc, unde c e a doua consoana, de exemplu

            return comb_lista


        return genereaza_combinatiile(vechiul_testament) + genereaza_combinatiile(noul_testament)
    

    SHORTCUTS_LOOKUP = [
        "sec", "fer", "sf", "etc", "pt",
    ] + lookup_biblice() + [re.compile(r'[iIdD]+.*[hH]+[rR]*.*')]


    def is_word_shortcut(text):

        for s in SHORTCUTS_LOOKUP:
            if type(s) == str:
                if s.lower() == text.lower():
                    return True
                    
            elif isinstance(s, re.Pattern):
                if s.search(text.lower()) is not None:
                    return True
                
        return False

    def between_two_numbers(text, delimitator_pos):
        # Ne asigurăm că pos nu este la începutul sau sfârșitul textului
        if delimitator_pos == 0 or delimitator_pos == len(text) - 1:
            return False

        # Obținem caracterul la poziția specificată
        char = text[delimitator_pos]

        # Divizăm textul în două părți
        before = text[:delimitator_pos+1]
        after = text[delimitator_pos:]

        # Construim regex-urile pentru a căuta grupuri de cifre înainte și după caracter
        regex_before = r'\d+\s*{}\s*$'.format(re.escape(char))
        regex_after = r'^\s*{}\s*\d+'.format(re.escape(char))

        # Verificăm dacă regex-urile se potrivesc cu părțile de text
        return re.search(regex_before, before) is not None and re.search(regex_after, after) is not None
        

    def extract_words_before_char(text, char_pos, num_words=1, include_given_char = False):
        words = text[:char_pos + (1 if include_given_char else 0)].split()
        return ' '.join(words[-num_words:]) if len(words) >= num_words else ' '.join(words)
        
    def extract_words_after_char(text, char_pos, num_words=1, include_given_char = False):
        words = text[char_pos + (0 if include_given_char else 1):].split()
        return ' '.join(words[:num_words]) if len(words) >= num_words else ' '.join(words)

    def is_numbered_list_index(text):
        pattern = re.compile(r'\b(([0-9]+)|([ivxlcdmIVXLCDM]+))\.')
        return pattern.search(text) is not None
        
    def is_number(text):
        def is_float(s):
            try:
                float(s)
                return True
            except ValueError:
                return False
        def is_integer(s):
            return s.isdigit()

        return is_float(text) or is_integer(text)

    paragraphs = []
    i = 0
    start_p = 0
    in_numbered_list = False
    in_paranthesis_counter = 0 # daca presupun ca nu o sa fie paranteze imbricate (structuri recursive), atunci pot face doar un boolean cu False si True; daca vreau sa tin cont de imbricari, tin un contor, dar problema e daca cineva uita sa inchida o paranteza (suntem, totusi, intr-un text literal, nu cod de programare) - as putea sa tin cont de o lungime maxim a parantezei, si sa se "prescrie" inchiderea ei daca cumva s-a uitat, si probabil ca e bine sa fac asa, ca si in cazul celelalt ar trebui facuta asa ceva - de ce ? fiindca daca presupun ca o paranteza e deschisa, orice punct din interior nu va mai avea efect de paragraf, si mi se pare normal, dar e o regula exhaustiva, daunatoare daca nu e contextul corect)
    in_paranthesis_positions = []
    in_paranthesis_max_length = 400
    def add_paragraph(to_p_inclusive):
        nonlocal paragraphs, start_p, i
        p = text[start_p:to_p_inclusive + 1].replace('\n', '').strip()
        if len(p) > 0:
            paragraphs.append(p)
            start_p = to_p_inclusive + 1
            return start_p + 1
        return i + 1
        
    def try_parse_paranthesis(current_char, current_position):
        nonlocal in_paranthesis_counter, in_paranthesis_positions
        if current_char == "(":
            in_paranthesis_counter += 1
            in_paranthesis_positions.append(current_position)
        elif current_char == ")":
            if in_paranthesis_counter >= 1:
                in_paranthesis_counter -= 1
                in_paranthesis_positions.pop()
            else:
                in_paranthesis_counter = 0
                in_paranthesis_positions = []
                print(f"AVERTISMENT: Avem o paranteza inchisa ) la pozitia {current_position} ({text[current_position-10:current_position+10]}) care nu are corespondent de paranteza deschisa - a fost ignorata")
        else:
            if in_paranthesis_counter > 0 and current_position - in_paranthesis_positions[-1] >= in_paranthesis_max_length:
                in_paranthesis_counter -= 1
                in_paranthesis_positions.pop()
                print(f"AVERTISMENT: ultima paranteza deschisa nu a mai avut corespondent dupa {in_paranthesis_max_length} caractere. AM inchis-o fortat")
        
    def acknowledge_i_position_change(new_i):
        if new_i < len(text):
            nonlocal i
            for p in range(i + 1, new_i):
                try_parse_paranthesis(text[p], p)
        else:
            print(f"AVERTISMENT: acknowledge_i_position_change: mi s-a dat o pozitie new_i care depaseste stringul initial: {new_i} (maximul e {len(text) - 1})")
        i = new_i
        
    while i < len(text):
        # luam delimitatoarele pe rand, si la lucruri din astea, e neaparat nevoie sa tratezi lucrurile cu o claritate si simplitate deosebita, fiindca sunt (inca) lucruri sintactice, dar de profunzime, si sintaxa inseamna structura in adancime

        # ceva logica de stare
        try_parse_paranthesis(text[i], i)

        # A. ! sau ? sunt delimitatoare sigure -> incheiam paragraful
        if text[i] == '!' or text[i] == '?':
            next_word = extract_words_after_char(text, i, 1)
            if next_word not in "!?:.": # daca urmeaza un alt delimitator, il ignoram pe asta
                acknowledge_i_position_change(add_paragraph(i))
                continue
            
        # B. punctul e punct (de delimitare paragraf) daca urmeaza dupa el o fraza. Are doar doua exceptii:
        #       - 1- scurtatura:
        #           -1.1- sub-exceptie, scurtatura de carti biblice, nu intotdeauna exista un standard dar se poate genera o multime de sub-scurtaturi pentru fiecare sursa biblica in parte - imi trebuie un lookup sau un regex sau o functie de validare
        #           -1.2- expresii "de epoca", adica cele din categoria inainte/dupa Mantuitorul Hristos (idH,iH,dH,...) - tot un lookup e bun si aici, sau un regex...
        #       - 2- index la o lista numerotata: grup de cifre sau litere (mici sau mari sau intercalate-greseala-tipizare) romane urmate de punctul nostru de aici - lista se recunoaste cel mai bine prin pastrarea starii (adica a analizelor anterioare) tinand seama ca am inceput o lista cand am avut un ":" anterior, avem acum sau am mai avut si dupa acest delimitator, indecsi de lista (descrisi mai sus) si inca nu am iesit din lista numerotata (iesire fiind un punct de delimitare care nu e urmat de un index de lista numerotata)
        
        elif text[i] == '.':
            next_word = extract_words_after_char(text, i, 1)
            if next_word not in "!?:.": # daca urmeaza un alt delimitator, il ignoram pe asta
            
                if in_paranthesis_counter <= 0: # NUMAI daca NU suntem in interiorul unor paranteze
                    
                    if not is_word_shortcut(extract_words_before_char(text, i, 1)) and not between_two_numbers(text, i):
                        # mai poate fi cazul 1 sau cazul cu "1." de exemplu (lista) - trebuie sa delimitam:
                        last_word = extract_words_before_char(text, i, 1, include_given_char=True)
                        if in_numbered_list and is_numbered_list_index(last_word):
                            #import pdb; pdb.set_trace()
                            pos_before_list_number = text.rfind(last_word)
                            if pos_before_list_number > -1:
                                acknowledge_i_position_change(add_paragraph(pos_before_list_number - 1))
                                continue
                            else:
                                print(f"avertisment: pos_before_list_number ar trebui sa fie mai mare ca -1 pentru last_word '{last_word}'")
                        else:
                            if not(is_number(extract_words_before_char(text, i, 1).strip()) and not in_numbered_list): # cazul versetelor, de exemplu, nu sunt in lista (nu am avut : ) dar sunt cu . dupa un numar, care e verset biblic sau orice altceva
                                if not is_numbered_list_index(extract_words_after_char(text, i, 1)):
                                    in_numbered_list = False
                                acknowledge_i_position_change(add_paragraph(i))
                                continue
            
        # C. ":" varianta de excludere standard ar fi aceea in care, ca la unele referentieri biblice nestandardizate, te pomenesti ca ai grup de cifre inainte si dupa acest semn
        #  si ca o optimizare, ideea ar fi sa nu inchei paragraful neaparat decat daca urmeaza o lista numerotata
        elif text[i] == ":":
            next_word = extract_words_after_char(text, i, 1)
            if next_word not in "!?:.": # daca urmeaza un alt delimitator, il ignoram pe asta

                if in_paranthesis_counter <= 0: # NUMAI daca NU suntem in interiorul unor paranteze
                
                    if is_numbered_list_index(extract_words_after_char(text, i, 1)) and not between_two_numbers(text, i):
                        in_numbered_list = True
                        acknowledge_i_position_change(add_paragraph(i))
                        continue
            
        # D. la ; o idee e sa nu fiu intre doua paranteze ( si ) ca atunci e clar ca nu inchei fraza;
        # iarasi cazul de excludere ca delimitator dintre doua grupuri de cifre
        # iarasi, sa nu fie intre doua numere, ca uneori, de ex la versete biblice, apare si acest delimitator
        # pentru economie de spatiu, daca nu ma aflu intr-o lista numerotata, atunci nu mai adaug un alt paragraf, ci las sa curga aceasta sectiune cu ; in continuare 
        # PE SCURT: NUMAI in liste numerotate, aplicam sectionare pe ; pentru paragraf nou
        elif text[i] == ';':
            next_word = extract_words_after_char(text, i, 1)
            if next_word not in "!?:.": # daca urmeaza un alt delimitator, il ignoram pe asta
                if in_paranthesis_counter <= 0: # NUMAI daca NU suntem in interiorul unor paranteze            
                    if not text[start_p:i].rfind("(") > text[start_p:i].rfind(")") and not between_two_numbers(text, i) and in_numbered_list:
                        acknowledge_i_position_change(add_paragraph(i))
                        continue
            
        i += 1 # pt cazurile in care nu s-a apelat add_paragraph si continue, incrementam pentru urmatorul caracter
        
    if len(text[start_p:i].strip()) > 0:
        paragraphs.append(text[start_p:i].replace('\n', '').strip())
    
    return paragraphs


def text_to_pdf(text, output_file, content_mode="html", split_mode="advanced"):
    doc = SimpleDocTemplate(output_file, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Divizăm textul în paragrafe
    paragraphs = split_in_paragraphs(text) if split_mode.lower() == "advanced" else text.split("\n")
    
    # Adăugăm fiecare paragraf în document
    flowables = []
    for paragraph in paragraphs:
        if content_mode.lower() == "html":
            soup = BeautifulSoup(paragraph, "html.parser")
            cleaned_text = soup.get_text()
            cleaned_text = cleaned_text.replace('<', ' ').replace('>', ' ')
            p = Paragraph(cleaned_text, styles["Normal"])
            flowables.append(p)
        else:
            cleaned_text = paragraph.replace('<', ' ').replace('>', ' ')
            if len(cleaned_text.strip()) > 0:
                flowables.append(Paragraph(cleaned_text, styles["Normal"]))
            else: 
                flowables.append(Spacer(1, 10))
    
    doc.build(flowables)
    #print(f"Fisierul PDF {output_file} a fost creat cu succes!")
    

def text_to_docx(text, output_file):
    # Creăm un document nou
    doc = Document()
    
    # Divizăm textul în paragrafe
    paragraphs = split_in_paragraphs(text)

    # Adăugăm fiecare paragraf în document
    for paragraph in paragraphs:
        soup = BeautifulSoup(paragraph, "html.parser")
        cleaned_text = soup.get_text()
        cleaned_text = cleaned_text.replace('<', ' ').replace('>', ' ')
        
        doc.add_paragraph(cleaned_text)

    # Salvăm documentul
    doc.save(output_file)
    print("Fișierul Word a fost creat cu succes!")
    
import unicodedata
def replace_special_characters(text):
    normalized_text = unicodedata.normalize('NFKD', text)
    replaced_text = normalized_text.encode('ASCII', 'ignore').decode('ASCII')
    return replaced_text
    
def normalize_romanian(text):
    # Conversia la lowercase
    text = text.lower()

    # Înlocuirea diacriticelor
    replacements = {
        '\u0103': 'a', '\u0102': 'a',  # ă
        '\u00E2': 'a', '\u00C2': 'a',  # â
        '\u00EE': 'i', '\u00CE': 'i',  # î
        '\u0219': 's', '\u0218': 's',  # ș
        '\u021B': 't', '\u021A': 't',  # ț
        '\u015F': 's', '\u015E': 's',  # ș cu cedilla
        '\u0163': 't', '\u0162': 't'   # ț cu cedilla
    }

    for diacritic, replacement in replacements.items():
        text = text.replace(diacritic, replacement)

    return text

import PyPDF2
from docx import Document

def read_pdf(file_path):
    pdf_file = open(file_path, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    pdf_text = ""
    for page_num in range(len(pdf_reader.pages)):
        page_text = pdf_reader.pages[page_num].extract_text()
        pdf_text += page_text + "\n"

    pdf_file.close()

    return pdf_text

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return ' '.join(full_text)